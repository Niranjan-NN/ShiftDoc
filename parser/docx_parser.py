"""
parser/docx_parser.py
---------------------
Handles extraction of content from .docx files using python-docx.
Returns structured data: headings, paragraphs, word count, etc.
"""

from docx import Document
from docx.oxml.ns import qn
import re


def parse_docx(file) -> dict:
    """
    Parse a .docx file and extract structured content.

    Args:
        file: A file-like object (e.g., from Streamlit's file_uploader)

    Returns:
        dict with keys: headings, paragraphs, tables, raw_text, word_count,
                        char_count, page_estimate, heading_count, para_count
    """
    try:
        doc = Document(file)
    except Exception as e:
        return {"error": f"Failed to open DOCX file: {str(e)}"}

    headings = []
    paragraphs = []
    raw_text_parts = []
    tables_data = []

    # ── Walk through every paragraph in the document ──────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue  # skip empty lines

        raw_text_parts.append(text)

        # Heading styles are named "Heading 1", "Heading 2", etc.
        style_name = ""

        if para.style is not None and para.style.name is not None:
            style_name = para.style.name

        if style_name.startswith("Heading"):
            headings.append({
                "level": style_name,
                "text": text
            })
        else:
            paragraphs.append(text)

    # ── Extract tables (bonus feature) ────────────────────────────────────
    for table_idx, table in enumerate(doc.tables):
        table_rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_rows.append(row_data)
        tables_data.append({
            "table_index": table_idx + 1,
            "rows": table_rows
        })

    raw_text = "\n".join(raw_text_parts)
    words = raw_text.split()
    word_count = len(words)
    char_count = len(raw_text)

    # DOCX files don't have a native "page count"; we estimate based on word count.
    # A standard page holds roughly 300 words.
    page_estimate = max(1, round(word_count / 300))

    return {
        "headings": headings,
        "paragraphs": paragraphs,
        "tables": tables_data,
        "raw_text": raw_text,
        "word_count": word_count,
        "char_count": char_count,
        "page_estimate": page_estimate,
        "heading_count": len(headings),
        "para_count": len(paragraphs),
        "error": None
    }